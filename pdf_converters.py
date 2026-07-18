import os
import fitz  # PyMuPDF
from PIL import Image
import docx
from openpyxl import load_workbook

def images_to_pdf(image_paths, output_path):
    """تحويل مجموعة من الصور (JPG/PNG) إلى ملف PDF موحد عالي الوضوح"""
    if not image_paths:
        raise ValueError("قائمة الصور فارغة.")
    doc = fitz.open()
    for img_path in image_paths:
        img = Image.open(img_path)
        # تحويل صيغة RGBA إلى RGB إن وجدت لتجنب الأخطاء
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        temp_jpg = img_path + "_converted.jpg"
        img.save(temp_jpg, "JPEG", quality=95)
        
        # إنشاء صفحة بنفس مقاس الصورة
        img_doc = fitz.open(temp_jpg)
        pdf_bytes = img_doc.convert_to_pdf()
        img_doc.close()
        
        pdf_page = fitz.open("pdf", pdf_bytes)
        doc.insert_pdf(pdf_page)
        pdf_page.close()
        
        if os.path.exists(temp_jpg):
            os.remove(temp_jpg)
            
    doc.save(output_path, garbage=4, deflate=True)
    doc.close()
    return output_path

def pdf_to_images(pdf_path, output_dir, dpi=150):
    """تحويل كل صفحة من صفحات الـ PDF إلى صورة JPG عالية الدقة"""
    doc = fitz.open(pdf_path)
    output_files = []
    zoom = dpi / 72.0
    mat = fitz.Matrix(zoom, zoom)
    for idx, page in enumerate(doc, start=1):
        pix = page.get_pixmap(matrix=mat)
        img_filename = os.path.join(output_dir, f"page_{idx}.jpg")
        pix.save(img_filename)
        output_files.append(img_filename)
    doc.close()
    return output_files

def pdf_to_word_text(pdf_path, output_docx_path):
    """تحويل محتوى PDF إلى مستند Word (.docx) قابل للتعديل بسهولة"""
    doc = fitz.open(pdf_path)
    word_doc = docx.Document()
    word_doc.add_heading("مستند محوّل من PDF", 0)
    
    for idx, page in enumerate(doc, start=1):
        text = page.get_text()
        if text.strip():
            p = word_doc.add_paragraph()
            p.add_run(f"--- صفحة {idx} ---\n").bold = True
            word_doc.add_paragraph(text)
    doc.close()
    word_doc.save(output_docx_path)
    return output_docx_path

def word_to_pdf(docx_path, output_pdf_path):
    """تحويل مستند Word (.docx) أو نص إلى PDF وثائقي منظم"""
    word_doc = docx.Document(docx_path)
    doc = fitz.open()
    page = doc.new_page()
    point = fitz.Point(50, 50)
    
    full_text = []
    for para in word_doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
            
    # كتابة النص داخل صفحة أو صفحات الـ PDF باستخدام خط قياسي أو عربي
    combined_text = "\n\n".join(full_text)
    # استخدام insert_textbox لتنسيق الهوامش والالتفاف التلقائي
    rect = fitz.Rect(50, 50, page.rect.width - 50, page.rect.height - 50)
    page.insert_textbox(rect, combined_text, fontsize=12, color=(0.1, 0.1, 0.1))
    
    doc.save(output_pdf_path, garbage=4, deflate=True)
    doc.close()
    return output_pdf_path

def excel_to_pdf(xlsx_path, output_pdf_path):
    """تحويل جدول Excel إلى مستند PDF منسق وسهل القراءة"""
    wb = load_workbook(xlsx_path, data_only=True)
    doc = fitz.open()
    page = doc.new_page()
    
    y_offset = 50
    rect = fitz.Rect(40, y_offset, page.rect.width - 40, page.rect.height - 40)
    
    table_lines = []
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        table_lines.append(f"=== جدول: {sheet_name} ===")
        for row in sheet.iter_rows(values_only=True):
            row_str = " | ".join([str(val) if val is not None else "" for val in row[:6]])
            if row_str.strip():
                table_lines.append(row_str)
        table_lines.append("\n")
        
    combined_text = "\n".join(table_lines)
    page.insert_textbox(rect, combined_text, fontsize=10, color=(0.1, 0.2, 0.4))
    doc.save(output_pdf_path)
    doc.close()
    return output_pdf_path
